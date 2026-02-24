from fastapi import APIRouter, Depends
from ..services.generation_service import generation_service
from pydantic import BaseModel
from typing import List
import random

router = APIRouter()

def extract_pdf_smart_chunks(reader, chunk_size: int = 4000, num_chunks: int = 3) -> str:
    """
    Multi-Chunk PDF Sampling (Option B):
    Divides the PDF into equal sections (beginning, middle, end),
    picks random pages from each section, and combines the text.
    This ensures every generation covers DIFFERENT topics from across
    the whole document — unique and varied each time.
    """
    total_pages = len(reader.pages)

    # Small document — just read it all, no need to sample
    if total_pages <= num_chunks * 2:
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text[:chunk_size * num_chunks]

    section_labels = ["Beginning", "Middle", "End"]
    collected = []
    section_size = total_pages // num_chunks

    for i in range(num_chunks):
        section_start = i * section_size
        section_end = min((i + 1) * section_size - 1, total_pages - 1)

        # Pick a random starting page within this section
        start_page = random.randint(section_start, max(section_start, section_end - 2))

        # Read a few pages from that starting point
        chunk_text = ""
        for page_idx in range(start_page, min(start_page + 4, total_pages)):
            chunk_text += reader.pages[page_idx].extract_text() or ""
            if len(chunk_text) >= chunk_size:
                break

        if chunk_text.strip():
            label = section_labels[i] if i < len(section_labels) else f"Section {i+1}"
            collected.append(f"[Document Section — {label} (pages ~{start_page+1}+)]\n{chunk_text[:chunk_size]}")

    return "\n\n---\n\n".join(collected)

class GenerateRequest(BaseModel):
    subject_name: str
    topic_name: str
    blooms_level: str
    subject_id: str = None
    count: int = 5
    rubric: dict = None
    engine: str = "local"
    custom_prompt: str = None
    fresh: bool = False
    kb_id: int = None

class BulkSaveRequest(BaseModel):
    subject_name: str
    topic_name: str
    questions: List[dict]
    marks: int = None
    duration: int = 60

from sqlalchemy.orm import Session
from ..database import get_db
from ..models import Question, Topic, Subject

@router.post("/questions")
def generate_questions(request: GenerateRequest, db: Session = Depends(get_db)):
    try:
        # 1. Generate via Unified AI Service (handles Local/Cloud/Fallback automatically)
        generated_data = generation_service.generate_questions(
            request.subject_name,
            request.topic_name,
            request.blooms_level,
            request.count,
            request.subject_id,
            request.rubric,
            request.engine,
            request.custom_prompt,
            request.fresh
        )

        # 2. Add temporary IDs for frontend pairing
        import uuid
        for q in generated_data:
            q['id'] = str(uuid.uuid4())
            
        return generated_data

    except Exception as e:
        import traceback
        error_msg = traceback.format_exc()
        print(f"Server Error: {error_msg}")
        # Return error as response for debugging (in prod setup this is bad, but fine for local debugging)
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=500, content={"error": "Internal Server Error", "traceback": error_msg})


from fastapi import UploadFile, File, Form
import shutil
import os

@router.post("/rubric/{rubric_id}")
async def generate_from_rubric(
    rubric_id: int, 
    engine: str = Form("local"),
    file: UploadFile = File(None),
    custom_prompt: str = Form(None),
    fresh: bool = Form(False),
    db: Session = Depends(get_db)
):
    """
    Generate exam questions based on a saved rubric
    Applies all rubric constraints including question types and LO distribution
    Optionally includes file context and custom prompt instructions.
    """
    from ..services.generation_service import generation_service
    
    context_text = None
    if file:
        try:
            os.makedirs("temp", exist_ok=True)
            file_path = f"temp/rubric_{file.filename}"
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            # Extract text using Multi-Chunk Sampling for variety
            _, ext = os.path.splitext(file_path)

            if ext.lower() == '.pdf':
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = extract_pdf_smart_chunks(reader, chunk_size=4000, num_chunks=3)
                print(f"[File] Smart-sampled PDF: {len(reader.pages)} pages → {len(text)} chars from 3 random sections")
            elif ext.lower() == '.docx':
                from docx import Document
                doc = Document(file_path)
                raw = "\n".join(p.text for p in doc.paragraphs)
                text = raw[:12000]
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read(12000)

            os.remove(file_path)
            if text.strip():
                context_text = text
        except Exception as e:
            print(f"File extraction error in rubric gen: {e}")

    try:
        result = generation_service.generate_from_rubric(
            rubric_id, 
            db, 
            engine=engine,
            context_text=context_text,
            custom_prompt=custom_prompt,
            fresh=fresh
        )
        
        return {
            "success": result["success"],
            "questions_generated": result["questions_generated"],
            "questions": result.get("questions", []),
            "log":  result["log"],
            "message": f"Successfully generated {result['questions_generated']} questions"
        }
    except ValueError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "error": f"Generation failed: {str(e)}"}

from fastapi import UploadFile, File, Form
import shutil
import os

@router.post("/from-file")
async def generate_from_file(
    file: UploadFile = File(...),
    count: int = Form(5),
    complexity: str = Form("Balanced"),
    engine: str = Form("local"),
    subject_id: str = Form(None),
    topic_id: str = Form(None),
    custom_prompt: str = Form(None),
    fresh: bool = Form(False),
    db: Session = Depends(get_db)
):
    """
    Generate questions directly from an uploaded file context, linked to a specific subject/topic.
    Includes memory-safe streaming extraction for massive PDFs.
    """
    try:
        # Save file temporarily
        os.makedirs("temp", exist_ok=True)
        file_path = f"temp/adhoc_{file.filename}"
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text using Multi-Chunk Sampling for variety across full document
        _, ext = os.path.splitext(file_path)

        if ext.lower() == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = extract_pdf_smart_chunks(reader, chunk_size=4000, num_chunks=3)
            print(f"[File] Smart-sampled PDF: {len(reader.pages)} pages → {len(text)} chars from 3 random sections")
        elif ext.lower() == '.docx':
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)[:12000]
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read(12000)

        # Cleanup temp file
        os.remove(file_path)
        
        if not text.strip():
            return {"error": "Could not extract text from file"}

        # Resolve Subject & Topic
        subject = None
        if subject_id:
            if str(subject_id).isdigit():
                subject = db.query(Subject).filter(Subject.id == int(subject_id)).first()
            else:
                subject = db.query(Subject).filter(Subject.code == subject_id).first()
        
        if not subject:
            # Fallback to "Uploaded Content" subject
            subject_name = "Uploaded Content"
            subject = db.query(Subject).filter(Subject.name == subject_name).first()
            if not subject:
                subject = Subject(
                    name=subject_name, 
                    code="FILE", 
                    color="#50FA7B", 
                    gradient="from-[#50FA7B] to-[#0A1F1F]",
                    introduction="Content generated from uploaded files"
                )
                db.add(subject)
                db.commit()
                db.refresh(subject)

        # Resolve Topic
        topic = None
        if topic_id:
             topic = db.query(Topic).filter(Topic.id == int(topic_id)).first()
        
        if not topic:
            topic_name = f"File: {file.filename}"
            topic = db.query(Topic).filter(Topic.subject_id == subject.id, Topic.name == topic_name).first()
            if not topic:
                topic = Topic(subject_id=subject.id, name=topic_name, has_syllabus=False)
                db.add(topic)
                db.commit()
                db.refresh(topic)

        # Use Unified Generation Service
        generated_data = generation_service.generate_questions_from_text(
            context_text=text,  # Context is already sampled/limited to ~12k chars internal to the logic
            subject_name=subject.name,
            topic_name=topic.name,
            count=count,
            complexity=complexity,
            engine=engine,
            custom_prompt=custom_prompt,
            fresh=fresh
        )
        
        # Add temporary IDs for frontend pairing
        import uuid
        for q in generated_data:
            q['id'] = str(uuid.uuid4())

        return generated_data

    except Exception as e:
        import traceback
        error_msg = traceback.format_exc()
        print(f"File Generation Error: {error_msg}")
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=500, content={"error": str(e), "trace": error_msg})

@router.post("/bulk-save")
def bulk_save_questions(request: BulkSaveRequest, db: Session = Depends(get_db)):
    """
    Saves a batch of approved questions and creates an exam history record.
    This is called AFTER vetting to ensure only approved content is persisted.
    """
    try:
        from ..models import Subject, Topic, Question, ExamHistory
        from ..services.logging_service import logging_service
        
        # 1. Resolve Subject
        subject = db.query(Subject).filter(Subject.name == request.subject_name).first()
        if not subject:
            # Check if auto-generated code already exists
            new_code = request.subject_name[:5].upper()
            existing_code = db.query(Subject).filter(Subject.code == new_code).first()
            if existing_code:
                # Append a random string or let it be longer since it's an edge case 
                import uuid
                new_code = (request.subject_name[:3] + str(uuid.uuid4())[:2]).upper()
                
            subject = Subject(
                name=request.subject_name,
                code=new_code,
                color="#50FA7B",
                gradient="from-[#50FA7B] to-[#0A1F1F]"
            )
            db.add(subject)
            db.commit()
            db.refresh(subject)
            
        # 2. Resolve Topic
        topic = db.query(Topic).filter(Topic.subject_id == subject.id, Topic.name == request.topic_name).first()
        if not topic:
            topic = Topic(subject_id=subject.id, name=request.topic_name, has_syllabus=False)
            db.add(topic)
            db.commit()
            db.refresh(topic)
            
        # 2.5 If no questions approved, just clear drafts and exit
        if not request.questions:
            db.query(Question).filter(Question.status == "draft").delete()
            db.commit()
            return {"success": True, "message": "All questions rejected. Cleared drafts."}
            
        # 3. Save Questions
        stored_questions_data = []
        for q_data in request.questions:
            new_q = Question(
                topic_id=topic.id,
                rubric_id=q_data.get('rubric_id'),
                question_text=q_data.get('question_text') or q_data.get('question', ''),
                question_type=q_data.get('question_type') or q_data.get('type') or 'MCQ',
                options=q_data.get('options', []),
                correct_answer=str(q_data.get('correct_answer', '')),
                explanation=q_data.get('explanation', ''),
                bloom_level=q_data.get('bloom_level') or q_data.get('complexity', 'Apply'),
                course_outcomes=q_data.get('courseOutcomes') or q_data.get('course_outcomes', {}),
                status="approved"
            )
            db.add(new_q)
            stored_questions_data.append({
                "question_text": new_q.question_text,
                "type": new_q.question_type,
                "marks": q_data.get('marks', 5)
            })
            
        db.commit()
        
        # 4. Create History
        total_marks = sum(q.get('marks', 5) for q in request.questions)
        new_history = ExamHistory(
            subject_name=subject.name,
            topic_name=request.topic_name,
            questions_count=len(request.questions),
            marks=total_marks,
            duration=request.duration,
            questions=request.questions
        )
        db.add(new_history)
        
        # 5. Log Activity
        logging_service.log_activity(db, "Exam Vetted & Saved", details={"subject": subject.name, "count": len(request.questions)})
        
        # 6. Cleanup any remaining drafts globally to prevent them from getting stuck in the Vetting Center
        db.query(Question).filter(Question.status == "draft").delete()
        db.commit()
        
        return {"success": True, "message": f"Saved {len(request.questions)} questions to history and cleared drafts."}
        
    except Exception as e:
        db.rollback()
        import traceback
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=500, content={"error": str(e), "trace": traceback.format_exc()})
